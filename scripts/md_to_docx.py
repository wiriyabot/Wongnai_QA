"""
แปลง docs/assignment_report_th.md → docs/assignment_report_th.docx
ใช้: uv run python scripts/md_to_docx.py
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

# ─── Paths ────────────────────────────────────────────────────────────────────
BASE   = Path(__file__).resolve().parent.parent
IN_MD  = BASE / "docs" / "assignment_report_th.md"
OUT_DX = BASE / "docs" / "assignment_report_th.docx"

# ─── Colour palette ───────────────────────────────────────────────────────────
C_DARK    = RGBColor(0x1A, 0x1A, 0x2E)   # h1/h2 background
C_H1_TXT  = RGBColor(0xFF, 0xFF, 0xFF)   # h1 text (white)
C_H2_TXT  = RGBColor(0x1A, 0x1A, 0x2E)  # h2 text (dark)
C_H3_TXT  = RGBColor(0x39, 0x9B, 0xD9)  # h3 text (blue accent)
C_CODE    = RGBColor(0x1E, 0x2D, 0x4A)  # inline code bg
C_ACCENT  = RGBColor(0xE9, 0x4F, 0x37)  # accent red
C_BODY    = RGBColor(0x1A, 0x1A, 0x2E)  # body text


def set_para_shading(para, hex_color: str):
    """ใส่สีพื้นหลัง paragraph"""
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def set_run_shading(run, hex_color: str):
    """ใส่ highlight สีให้ run (inline code)"""
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    rPr.append(shd)


def add_h1(doc: Document, text: str):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    set_para_shading(p, "1A1A2E")
    run = p.add_run(text)
    run.font.color.rgb = C_H1_TXT
    run.font.bold = True
    run.font.size = Pt(18)
    run.font.name = "Sarabun"
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_h2(doc: Document, text: str):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    run = p.add_run(text)
    run.font.color.rgb = C_H2_TXT
    run.font.bold = True
    run.font.size = Pt(15)
    run.font.name = "Sarabun"
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    # left border accent
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"),  "20")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "E94F37")
    pBdr.append(left)
    pPr.append(pBdr)
    return p


def add_h3(doc: Document, text: str):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 3"]
    run = p.add_run(text)
    run.font.color.rgb = C_H3_TXT
    run.font.bold = True
    run.font.italic = False
    run.font.size = Pt(13)
    run.font.name = "Sarabun"
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_body(doc: Document, text: str, indent=0):
    """เพิ่ม paragraph ธรรมดา รองรับ **bold** และ `code` inline"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(indent * 0.25)
    _add_inline(p, text)
    return p


def _add_inline(para, text: str):
    """แยก bold / code / ปกติ และเพิ่มลงใน paragraph"""
    # Pattern: **bold**, *italic*, `code`
    token_re = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\$.*?\$)')
    parts = token_re.split(text)
    for part in parts:
        if not part:
            continue
        run = para.add_run()
        run.font.name = "Sarabun"
        run.font.size = Pt(12)
        run.font.color.rgb = C_BODY
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.font.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]
            run.font.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Courier New"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0xE9, 0x4F, 0x37)
            set_run_shading(run, "EEF2FF")
        elif part.startswith("$") and part.endswith("$"):
            # Inline math — render as italic
            run.text = part[1:-1]
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x1A, 0x82, 0xFF)
        else:
            run.text = part


def add_bullet(doc: Document, text: str, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.2)
    _add_inline(p, text)


def add_numbered(doc: Document, text: str, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.2)
    _add_inline(p, text)


def add_code_block(doc: Document, lines: list[str]):
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name  = "Courier New"
        run.font.size  = Pt(10)
        run.font.color.rgb = RGBColor(0xE9, 0x4F, 0x37)
        set_para_shading(p, "EEF2FF")
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Inches(0.3)


def add_table(doc: Document, rows: list[list[str]]):
    """สร้างตาราง จาก rows ของ Markdown table"""
    if len(rows) < 2:
        return
    # rows[0] = header, rows[1] = separator, rows[2:] = data
    headers = [c.strip() for c in rows[0] if c.strip()]
    data    = [[c.strip() for c in row if c.strip()] for row in rows[2:]]

    tbl = doc.add_table(rows=1+len(data), cols=len(headers))
    tbl.style = "Table Grid"

    # Header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.name = "Sarabun"
        run.font.size = Pt(11)
        run.font.color.rgb = C_H1_TXT
        set_para_shading(cell.paragraphs[0], "1A1A2E")

    # Data rows
    for r_idx, row in enumerate(data):
        tbl_row = tbl.rows[r_idx + 1]
        bg = "FFFFFF" if r_idx % 2 == 0 else "F0F4FF"
        for c_idx, cell_text in enumerate(row):
            if c_idx < len(tbl_row.cells):
                cell = tbl_row.cells[c_idx]
                cell.text = ""
                _add_inline(cell.paragraphs[0], cell_text)
                set_para_shading(cell.paragraphs[0], bg)
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(11)

    doc.add_paragraph()  # gap after table


def add_blockquote(doc: Document, text: str):
    p = doc.add_paragraph()
    _add_inline(p, text)
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_after  = Pt(3)
    # left border blue
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"),  "16")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "399BD9")
    pBdr.append(left)
    pPr.append(pBdr)
    set_para_shading(p, "EEF5FF")


# ─── Main parser ──────────────────────────────────────────────────────────────

def convert(md_path: Path, out_path: Path):
    doc = Document()

    # Page margins
    for sect in doc.sections:
        sect.top_margin    = Inches(1.0)
        sect.bottom_margin = Inches(1.0)
        sect.left_margin   = Inches(1.15)
        sect.right_margin  = Inches(1.15)

    # Default body font
    doc.styles["Normal"].font.name = "Sarabun"
    doc.styles["Normal"].font.size = Pt(12)

    lines = md_path.read_text(encoding="utf-8").splitlines()

    i = 0
    table_buf: list[list[str]] = []
    code_buf:  list[str]       = []
    in_code  = False
    code_lang = ""

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        # ── Code fence ────────────────────────────────────────────────────────
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = stripped[3:].strip()
                code_buf  = []
            else:
                # end fence
                in_code = False
                if code_lang.lower() not in ("mermaid",):  # skip mermaid rendering
                    note = f"[Diagram: {code_lang}]" if code_lang == "mermaid" else None
                    if note:
                        add_body(doc, f"📊 {note} (แสดงใน Markdown viewer)")
                    else:
                        add_code_block(doc, code_buf)
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(raw)
            i += 1
            continue

        # ── Table ─────────────────────────────────────────────────────────────
        if stripped.startswith("|"):
            row_cells = [c for c in stripped.split("|")]
            table_buf.append(row_cells)
            i += 1
            continue
        else:
            if table_buf:
                add_table(doc, table_buf)
                table_buf = []

        # ── Headings ──────────────────────────────────────────────────────────
        if stripped.startswith("#### "):
            add_h3(doc, stripped[5:])
        elif stripped.startswith("### "):
            add_h3(doc, stripped[4:])
        elif stripped.startswith("## "):
            add_h2(doc, stripped[3:])
        elif stripped.startswith("# "):
            add_h1(doc, stripped[2:])

        # ── Ordered list ──────────────────────────────────────────────────────
        elif re.match(r"^\d+\.\s", stripped):
            text  = re.sub(r"^\d+\.\s*", "", stripped)
            level = (len(raw) - len(raw.lstrip())) // 3
            add_numbered(doc, text, level=level)

        # ── Unordered list ────────────────────────────────────────────────────
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text  = stripped[2:]
            level = (len(raw) - len(raw.lstrip())) // 2
            add_bullet(doc, text, level=level)

        # ── Blockquote ────────────────────────────────────────────────────────
        elif stripped.startswith("> "):
            add_blockquote(doc, stripped[2:])

        # ── Horizontal rule ───────────────────────────────────────────────────
        elif stripped in ("---", "===", "***"):
            p = doc.add_paragraph()
            pPr  = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot  = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single")
            bot.set(qn("w:sz"), "6")
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), "E94F37")
            pBdr.append(bot)
            pPr.append(pBdr)

        # ── Blank line ────────────────────────────────────────────────────────
        elif stripped == "":
            pass  # skip blanks silently

        # ── Normal paragraph ──────────────────────────────────────────────────
        else:
            add_body(doc, stripped)

        i += 1

    # Flush remaining table
    if table_buf:
        add_table(doc, table_buf)

    doc.save(str(out_path))
    print(f"[OK] Saved: {out_path}")


if __name__ == "__main__":
    convert(IN_MD, OUT_DX)
